import xml.etree.ElementTree as ET
import base64
from typing import List, Dict
import argparse
from tabulate import tabulate
import os, re, json
import tempfile
import subprocess
from html2docx import html2docx
from collections import defaultdict
from docxcompose.composer import Composer
from docx import Document as Document_compose
from docxtpl import *
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# Common Functions
def clean_text(text):
    c_text = re.sub(r' +', ' ', text)
    c_text = c_text.replace("\n", "")
    return text

def check_integer(value):
    try:
        if not isinstance(value, int):
            raise ValueError("Error: The provided value is not an integer.")
        return f"The value {value} is a valid integer."
    except ValueError as e:
        return str(e)
def check_string(value):
    try:
        if not isinstance(value, str):
            raise ValueError("Error: The provided value is not an integer.")
        return f"The value {value} is a valid string."
    except ValueError as e:
        return str(e)

class SecurityReportGenerator:
    def __init__(self):
        self.severity_colors = {
            'High': 'FF0000',      # Red
            'Medium': 'FFA500',    # Orange
            'Low': 'FFFF00',       # Yellow
            'Information': '0000FF' # Blue
        }       

    def create_shaded_cell(self, cell, color):
        """Apply background shading to a table cell"""
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm)

    def setup_styles(self, document):
        try:
            # Get or create styles
            if not hasattr(document, 'styles'):
                document._part.load_style_definitions()
            
            styles = document.styles

            # Base style
            try:
                style = styles["Normal"]
            except KeyError:
                style = styles.add_style("Normal", WD_STYLE_TYPE.PARAGRAPH)
            font = style.font
            font.name = "Calibri Light"
            font.size = Pt(10)
            
            # Finding title style
            if 'FindingTitle' not in styles:
                title_style = styles.add_style('FindingTitle', WD_STYLE_TYPE.PARAGRAPH)
                title_style.font.name = "Calibri"
                title_style.font.size = Pt(16)
                title_style.font.bold = True
                title_style.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
            
            # Section heading style
            if 'SectionHeading' not in styles:
                section_style = styles.add_style('SectionHeading', WD_STYLE_TYPE.PARAGRAPH)
                section_style.font.name = "Calibri"
                section_style.font.size = Pt(12)
                section_style.font.bold = True
                section_style.font.color.rgb = RGBColor(51, 51, 51)  # Dark gray
                
            # Subheading style
            if 'Subheading' not in styles:
                subheading_style = styles.add_style('Subheading', WD_STYLE_TYPE.PARAGRAPH)
                subheading_style.font.name = "Calibri"
                subheading_style.font.size = Pt(11)
                subheading_style.font.bold = True
                subheading_style.font.underline = True
                subheading_style.font.color.rgb = RGBColor(51, 51, 51)  # Dark gray
            
            # Code block style
            if 'CodeBlock' not in styles:
                code_style = styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
                code_style.font.name = "Consolas"
                code_style.font.size = Pt(9)
                code_style.font.color.rgb = RGBColor(51, 51, 51)
                code_style.paragraph_format.space_before = Pt(6)
                code_style.paragraph_format.space_after = Pt(6)
                code_style.paragraph_format.left_indent = Inches(0.3)

        except Exception as e:
            #print(f"Warning: Error setting up styles: {str(e)}")
            # Continue without styles rather than failing completely
            pass

    def format_issue_detail(self, doc, issue_detail):
        """Format the issue detail section with proper styling"""
        if issue_detail == 'N/A':
            return False
            
        # Split content by known subheadings
        sections = {}
        current_section = []
        current_heading = "main"
        
        for line in issue_detail.split('\n'):
            if line.strip().startswith("Affecting Versions:"):
                current_heading = "Affecting Versions"
                sections[current_heading] = []
            elif line.strip().startswith("Other Considerations:"):
                current_heading = "Other Considerations"
                sections[current_heading] = []
            else:
                if current_heading in sections:
                    sections[current_heading].append(line)
                else:
                    current_section.append(line)
        
        if current_section:
            sections["main"] = current_section
            
        # Add the main content
        if "main" in sections:
            main_text = '\n'.join(sections["main"]).strip()
            if main_text:
                p = doc.add_paragraph(main_text)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.space_after = Pt(12)
        
        # Add subheadings and their content
        for heading in ["Affecting Versions", "Other Considerations"]:
            if heading in sections and sections[heading]:
                # Add subheading
                p = doc.add_paragraph(heading + ":", style='Subheading')
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(6)
                
                # Add content under subheading
                content = '\n'.join(sections[heading]).strip()
                if content:
                    p = doc.add_paragraph(content)
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.space_after = Pt(12)
        
        return True

    def make_report(self, output_data: list):
        """Creates a complete security findings report"""
        target_doc = Document_compose("template/finding.docx")
        composer = Composer(target_doc)

        for finding in output_data:
            doc = DocxTemplate("template/finding.docx")
            self.setup_styles(doc)
            
            # Create basic context without N/A values
            context = {
                'finding_name': finding['name'],
                'finding_host': finding['host'],
                'finding_path': finding['path'],
                'finding_severity': finding['severity'],
                'finding_confidence': finding['confidence']
            }
            
            # Add optional fields only if they're not N/A
            if finding['description'] != 'N/A':
                context['issue_background'] = finding['description']
            
            if finding['remediation_details'] != 'N/A':
                context['issue_remediation'] = finding['remediation_details']
            
            if finding['vulnerabilityClassifications'] != 'N/A':
                context['vulnerability_classification'] = finding['vulnerabilityClassifications']
            
            # Handle affected locations
            if finding.get('affected_locations'):
                locations = [loc for loc in finding['affected_locations'] if loc != 'N/A']
                if locations:
                    locations_text = '\n'.join(f'• {loc}' for loc in locations)
                    context['issue_locations'] = locations_text
            
            doc.render(context)
            
            # Add Issue Detail section only if it's not N/A
            if finding.get('issueDetail') and finding['issueDetail'] != 'N/A':
                issue_detail_header = doc.add_paragraph('Issue Detail', style='SectionHeading')
                issue_detail_header.paragraph_format.space_before = Pt(12)
                issue_detail_header.paragraph_format.space_after = Pt(6)
                self.format_issue_detail(doc, finding['issueDetail'])
            
            # Add HTTP messages section only if there are valid messages
            if finding.get('sample_http_messages'):
                valid_messages = [msg for msg in finding['sample_http_messages'] 
                                if msg.get('request') != 'N/A' or msg.get('response') != 'N/A']
                
                if valid_messages:
                    http_section = doc.add_paragraph('Sample HTTP Messages\n')
                    http_section.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    for i, message in enumerate(valid_messages, 1):
                        if len(valid_messages) > 1:
                            msg_header = doc.add_paragraph(f'Message Pair #{i}', style='SectionHeading')
                            msg_header.paragraph_format.left_indent = Inches(0.2)
                        
                        if message.get('request') and message['request'] != 'N/A':
                            req_header = doc.add_paragraph('Request:', style='SectionHeading')
                            req_header.paragraph_format.left_indent = Inches(0.2)
                            req_content = doc.add_paragraph(message['request'])
                            req_content.paragraph_format.left_indent = Inches(0.4)
                        
                        if message.get('response') and message['response'] != 'N/A':
                            resp_header = doc.add_paragraph('Response:', style='SectionHeading')
                            resp_header.paragraph_format.left_indent = Inches(0.2)
                            resp_content = doc.add_paragraph(message['response'])
                            resp_content.paragraph_format.left_indent = Inches(0.4)
                        
                        if i < len(valid_messages):
                            separator = doc.add_paragraph('─' * 40)
                            separator.paragraph_format.left_indent = Inches(0.4)
            
            doc.add_page_break()
            composer.append(doc)

        composer.save('complete.docx')
        print("\nReport generated successfully as 'security_findings_report.docx'")

class BurpScanParser:
    def __init__(self, xml_file_path: str):
        """
        Initialize the parser with the Burp Suite XML export file
        """
        self.tree = ET.parse(xml_file_path)
        self.root = self.tree.getroot()
        self.raw_vulnerabilities = self.extract_raw_vulnerabilities()
        self.grouped_vulnerabilities = self.group_vulnerabilities()
    def _decode_base64(self, encoded_str: str) -> str:
        """
        Decode base64 encoded strings safely with better error handling
        """
        if not encoded_str:
            return ""
        
        try:
            # Remove any whitespace that might be in the base64 string
            cleaned_str = ''.join(encoded_str.split())
            # Add padding if necessary
            padding = 4 - (len(cleaned_str) % 4)
            if padding != 4:
                cleaned_str += '=' * padding
            
            decoded = base64.b64decode(cleaned_str)
            # Try UTF-8 first
            try:
                return decoded.decode('utf-8')
            except UnicodeDecodeError:
                # Fall back to latin-1 if UTF-8 fails
                return decoded.decode('latin-1')
        except Exception as e:
            return f"Decoding failed: {str(e)}"
    def _format_http_message(self, message: str) -> str:
        """
        Format HTTP message for better readability
        """
        if not message:
            return "No content"

        # Split headers and body
        parts = message.split('\r\n\r\n', 1)
        if len(parts) < 2:
            parts = message.split('\n\n', 1)
        
        if len(parts) == 2:
            headers, body = parts
        else:
            headers, body = message, ""

        # Format headers
        formatted_headers = []
        for line in headers.split('\n'):
            if ': ' in line:
                header_name, header_value = line.split(': ', 1)
                formatted_headers.append(f"{header_name}: {header_value}")
            else:
                formatted_headers.append(line)

        # Try to format body if it looks like JSON
        if body.strip().startswith('{') or body.strip().startswith('['):
            try:
                import json
                parsed_body = json.loads(body)
                body = json.dumps(parsed_body, indent=2)
            except:
                pass

        return '\n'.join(formatted_headers) + '\n\n' + body if body else '\n'.join(formatted_headers)
    def _format_vulnerability_classifications(self, raw_text: str) -> str:
        """
        Format vulnerability classification links from HTML to a clean, readable format.
        
        Args:
            raw_text (str): Raw HTML text containing vulnerability classifications
            
        Returns:
            str: Formatted vulnerability classifications text
        """
        if raw_text == 'N/A':
            return 'N/A'
        
        # Use regex to extract information from the HTML links
        pattern = r'<a href="(https?://[^"]+)">([^<]+)</a>'
        matches = re.findall(pattern, raw_text)
        
        if not matches:
            return raw_text
        
        # Format each classification on a new line
        formatted_lines = []
        for url, text in matches:
            formatted_lines.append(f"{text}")
            formatted_lines.append(f"Reference: {url}")
            formatted_lines.append("")  # Add blank line between entries
        
        return "\n".join(formatted_lines).rstrip()
    def extract_raw_vulnerabilities(self) -> List[Dict]:

        vulnerabilities = []
        
        for vuln in self.root.findall('.//issue'):
            vulnerability = {
                'name': vuln.findtext('name', 'N/A'),
                'severity': vuln.findtext('severity', 'N/A'),
                'confidence': vuln.findtext('confidence', 'N/A'),
                'host': vuln.findtext('host', 'N/A'),
                'path': vuln.findtext('path', 'N/A'),
                'description': vuln.findtext('issueBackground', 'N/A'),
                'issueDetail': vuln.findtext('issueDetail', 'N/A'),
                'remediation_details': vuln.findtext('remediationBackground', 'N/A'),
                'vulnerabilityClassifications': self._format_vulnerability_classifications(vuln.findtext('vulnerabilityClassifications', 'N/A')),
                'http_messages': []
            }
            
           # print(vulnerability)
            # Extract all HTTP messages
            for requestresponse in vuln.findall('.//requestresponse'):
                request_elem = requestresponse.find('request')
                response_elem = requestresponse.find('response')
                
                # Process request
                request_text = ''
                if request_elem is not None:
                    request_text = request_elem.text if request_elem.text else ''
                    # Check if base64 encoded
                    if request_elem.get('base64') == 'true':
                        request_text = self._decode_base64(request_text)
                    request_text = self._truncate_cookie(request_text)
                
                # Process response
                response_text = ''
                if response_elem is not None:
                    response_text = response_elem.text if response_elem.text else ''
                    # Check if base64 encoded
                    if response_elem.get('base64') == 'true':
                        response_text = self._decode_base64(response_text)
                
                # Add the message pair to the list
                message = {
                    'request': self._format_http_message(request_text) if request_text else 'No request data available',
                    'response': self._format_http_message(response_text) if response_text else 'No response data available'
                }
                
                vulnerability['http_messages'].append(message)
            
            vulnerabilities.append(vulnerability)
        
        return vulnerabilities        
    def _truncate_cookie(self, http_message: str) -> str:
        """
        Truncates only the Cookie header in the HTTP request while preserving all other headers
        and message content.
        
        Args:
            http_message (str): The complete HTTP message
            
        Returns:
            str: HTTP message with only Cookie header truncated
        """
        if not http_message:
            return http_message
            
        # Split the message into lines
        lines = http_message.split('\n')
        modified_lines = []
        
        for line in lines:
            # Only modify lines that start with 'Cookie:'
            if line.strip().startswith('Cookie:'):
                modified_lines.append('Cookie: [snipped]')
            else:
                modified_lines.append(line)
                
        # Reconstruct the message maintaining original line endings
        return '\n'.join(modified_lines)
    def _show_http_messages(self, instance: Dict):
        """
        Display all HTTP messages for a specific instance with improved formatting.
        Each instance may contain multiple request-response pairs.
        
        Args:
            instance (Dict): Instance dictionary containing HTTP messages
        """
        if not instance['http_messages']:
            print("\nNo HTTP messages available for this instance.")
            return

        for i, msg in enumerate(instance['http_messages'], 1):
            print(f"\n{'='*80}")
            print(f"HTTP Message Pair #{i}")
            print(f"{'='*80}")

            # Request
            print("\nREQUEST:")
            print("-" * 40)
            print(msg.get('request', 'No request data available'))

            # Response
            print("\nRESPONSE:")
            print("-" * 40)
            print(msg.get('response', 'No response data available'))
            
            # Add separator between message pairs
            if i < len(instance['http_messages']):
                print("\n" + "="*80)
                print("Next Message Pair")
                
        # Add final separator
        print("\n" + "="*80)
    def group_vulnerabilities(self) -> List[Dict]:
        """
        Group vulnerabilities by name and sort by severity
        """
        severity_order = {'Critical':0, 'High': 1, 'Medium': 2, 'Low': 3, 'Information': 4}
        grouped = defaultdict(list)
        
        for vuln in self.raw_vulnerabilities:
            key = (vuln['name'], vuln['severity'], vuln['confidence'])
            grouped[key].append(vuln)
        
        result = []
        for (name, severity, confidence), instances in grouped.items():
            result.append({
                'name': name,
                'severity': severity,
                'confidence': confidence,
                'instances': instances
            })
        
        return sorted(result, key=lambda x: (severity_order.get(x['severity'], 999), x['name']))
    def delete_finding(self, finding_input: str) -> bool:
        """
        Delete one or multiple findings from the grouped vulnerabilities
        
        Args:
            finding_input (str): Single number or comma-separated string of finding numbers to delete
            
        Returns:
            bool: True if all specified findings were deleted successfully, False otherwise
        """
        # Handle empty input
        if not finding_input.strip():
            print("Error: No finding number provided")
            return False
        
        # Parse and validate the input
        try:
            # Check if input contains comma (multiple numbers)
            if ',' in finding_input:
                # Split the input string and convert to integers
                numbers = [int(num.strip()) for num in finding_input.split(',')]
                # Remove duplicates while maintaining order
                numbers = list(dict.fromkeys(numbers))
            else:
                # Single number case
                numbers = [int(finding_input.strip())]
                
        except ValueError:
            if ',' in finding_input:
                print("Error: Invalid input format. Please provide comma-separated numbers (e.g., '1,2,3')")
            else:
                print("Error: Invalid input. Please provide a valid finding number")
            return False
        
        # Validate that all numbers are within range
        max_findings = len(self.grouped_vulnerabilities)
        invalid_numbers = [num for num in numbers if num < 1 or num > max_findings]
        
        if invalid_numbers:
            if len(invalid_numbers) == 1:
                print(f"Error: Finding #{invalid_numbers[0]} does not exist. Valid range is 1-{max_findings}.")
            else:
                print(f"Error: Finding numbers {', '.join(map(str, invalid_numbers))} do not exist. "
                      f"Valid range is 1-{max_findings}.")
            return False
        
        # Sort numbers in descending order to avoid index shifting when deleting
        numbers.sort(reverse=True)
        
        # Store deleted findings for confirmation message
        deleted_findings = []
        
        # Delete the findings
        for num in numbers:
            deleted_finding = self.grouped_vulnerabilities.pop(num - 1)
            deleted_findings.append(f"#{num} - {deleted_finding['name']}")
        
        # Print confirmation message
        if len(deleted_findings) == 1:
            print(f"\nDeleted finding: {deleted_findings[0]}")
        else:
            print("\nDeleted the following findings:")
            for finding in reversed(deleted_findings):  # Reverse to show in original order
                print(finding)
        
        return True
    def modify_finding(self, finding_number: int):
        """
        Modify details of a specific finding using a text editor
        """ 
        if not (1 <= finding_number <= len(self.grouped_vulnerabilities)):
            print(f"Error: Finding #{finding_number} does not exist.")
            return

        finding = self.grouped_vulnerabilities[finding_number - 1]
        instance = finding['instances'][0]  # Get first instance for editing

        # Create a dictionary with editable fields
        editable_data = {
            "name": finding['name'],
            "severity": finding['severity'],
            "confidence": finding['confidence'],
            "description": instance['description'],
            "remediation_details": instance['remediation_details'],
            "issueDetail": instance['issueDetail'],
            "vulnerabilityClassifications": instance['vulnerabilityClassifications']
        }

        # Add help text and format as JSON
        editor_content = {
            "INSTRUCTIONS": {
                "1": "Edit the values below to modify the finding",
                "2": "Severity must be one of: Critical,High, Medium, Low, Information",
                "3": "Confidence must be one of: Certain, Firm, Tentative, Informative",
                "4": "Save and exit the editor to apply changes",
                "5": "Exit without saving to cancel"
            },
            "FINDING_DATA": editable_data
        }

        # Create temporary file
        with tempfile.NamedTemporaryFile(mode='w', suffix='.json', delete=False) as temp_file:
            json.dump(editor_content, temp_file, indent=4)
            temp_filepath = temp_file.name

        try:
            # Open the default system editor
            if os.name == 'nt':  # Windows
                os.startfile(temp_filepath)
            else:  # Unix-like systems
                editor = os.environ.get('EDITOR', 'nano')  # Default to nano if EDITOR not set
                subprocess.call([editor, temp_filepath])

            # Read the modified content
            with open(temp_filepath, 'r') as temp_file:
                try:
                    modified_data = json.load(temp_file)
                    modified_finding = modified_data['FINDING_DATA']

                    # Validate severity and confidence
                    if modified_finding['severity'] not in ['Critical','High', 'Medium', 'Low', 'Information']:
                        raise ValueError("Invalid severity level")
                    if modified_finding['confidence'] not in ['Certain', 'Firm', 'Tentative']:
                        raise ValueError("Invalid confidence level")

                    # Update finding and all instances
                    finding['name'] = modified_finding['name']
                    finding['severity'] = modified_finding['severity']
                    finding['confidence'] = modified_finding['confidence']

                    for instance in finding['instances']:
                        instance['name'] = modified_finding['name']
                        instance['severity'] = modified_finding['severity']
                        instance['confidence'] = modified_finding['confidence']
                        instance['description'] = modified_finding['description']
                        instance['remediation_details'] = modified_finding['remediation_details']
                        instance['issueDetail'] = modified_finding['issueDetail']
                        instance['vulnerabilityClassifications'] = modified_finding['vulnerabilityClassifications']

                    # Regroup vulnerabilities to maintain proper sorting
                    self.grouped_vulnerabilities = self.group_vulnerabilities()
                    print("\nFinding updated successfully!")

                except json.JSONDecodeError:
                    print("\nError: Invalid JSON format in edited file. Changes not saved.")
                except ValueError as e:
                    print(f"\nError: {str(e)}. Changes not saved.")
                except KeyError as e:
                    print(f"\nError: Required field {str(e)} missing. Changes not saved.")

        finally:
            # Clean up temporary file
            try:
                os.unlink(temp_filepath)
            except:
                pass
    def save_to_json(self):
        """
        Save findings to a JSON file with selected information.
        Includes all HTTP messages for each finding instead of just the first one.
        """
        output_data = []
        
        for vuln in self.grouped_vulnerabilities:
            # Take the first instance as representative for general information
            instance = vuln['instances'][0]
            
            # Get all HTTP messages
            sample_http = []
            if instance['http_messages']:
                for http_message in instance['http_messages']:
                    message_pair = {
                        'request': http_message['request'],
                        'response': http_message['response']
                    }
                    sample_http.append(message_pair)
            
            # Create affected locations list
            affected_locations = [f"{inst['host']}{inst['path']}" for inst in vuln['instances']]
            
            finding_data = {
                'name': vuln['name'],
                'severity': vuln['severity'],
                'confidence': vuln['confidence'],
                'host': instance['confidence'],
                'path': instance['confidence'],
                'description': instance['description'],
                'remediation_details': instance['remediation_details'],
                'issueDetail': instance['issueDetail'],
                'affected_locations': affected_locations,
                'vulnerabilityClassifications': instance['vulnerabilityClassifications'],
                'sample_http_messages': sample_http  # Changed from sample_http_message to sample_http_messages
            }
            output_data.append(finding_data)

        return output_data

        print(f"\nFindings saved.")
    def display_table(self):
        """
        Display vulnerabilities in a tabulated format
        """
        table_data = []
        for i, vuln in enumerate(self.grouped_vulnerabilities, 1):
            table_data.append([
                i,
                vuln['name'],
                vuln['severity']
            ])
        
        headers = ['#', 'Vulnerability', 'Severity']
        print(tabulate(table_data, headers=headers, tablefmt='grid'))
    def show_detailed_finding(self, finding_number: int):
        """
        Display detailed information about a specific finding
        """
        if not (1 <= finding_number <= len(self.grouped_vulnerabilities)):
            print(f"Error: Finding #{finding_number} does not exist.")
            return

        vuln_group = self.grouped_vulnerabilities[finding_number - 1]
        
        print("\n" + "="*80)
        print(f"DETAILED INFORMATION FOR FINDING #{finding_number}")
        print("="*80 + "\n")
        
        # Basic Information
        print(f"Name: {vuln_group['name']}")
        print(f"Severity: {vuln_group['severity']}")
        print(f"Confidence: {vuln_group['confidence']}")
        
        if (vuln_group['instances'][0]['description'] == 'N/A' and vuln_group['instances'][0]['remediation_details'] == 'N/A'):
            print("\nIssueDetail:")
            print("-" * 40)
            print(vuln_group['instances'][0]['issueDetail'])
        else:
            # Description
            print("\nDESCRIPTION:")
            print("-" * 40)
            print(vuln_group['instances'][0]['description'])
            
            # Remediation
            print("\nREMEDIATION:")
            print("-" * 40)
            print(vuln_group['instances'][0]['remediation_details'])
        
        # Affected Locations
        print("\nAFFECTED LOCATIONS:")
        print("-" * 40)
        for i, instance in enumerate(vuln_group['instances'], 1):
            print(f"{i}. {instance['host']}{instance['path']}")

        # Vulnerability Classifications
        if vuln_group['instances'][0]['vulnerabilityClassifications'] != 'N/A':
            print("\nVULNERABILITY CLASSIFICATIONS:")
            print("-" * 40)
            print(vuln_group['instances'][0]['vulnerabilityClassifications'])

        
        while True:
            print("\nOptions:")
            print("1-N: View HTTP messages for a specific location")
            print("b: Go back to main menu")
            print("d: Delete this finding")
            choice = input("\nEnter your choice: ").lower()
            
            if choice == 'd':
                burp_parser.delete_finding(str(finding_number)) 
                break             
            elif choice == 'b':
                break
            
            try:
                instance_num = int(choice)
                if 1 <= instance_num <= len(vuln_group['instances']):
                    self._show_http_messages(vuln_group['instances'][instance_num - 1])
                    input("\nPress Enter to continue...")
                else:
                    print("Invalid location number.")
            except ValueError:
                print("Invalid input. Please enter a number or 'b'.")   

parser = argparse.ArgumentParser(description='Parse Burp Suite XML Scan Export')
parser.add_argument('xml_file', help='Path to Burp Suite XML export file')
args = parser.parse_args()

try:
    burp_parser = BurpScanParser(args.xml_file)
    
    while True:
        os.system('cls' if os.name == 'nt' else 'clear')  # Clear screen
        print(r"""

                ______                      _____         _  _         
                | ___ \                    /  ___|       (_)| |        
                | |_/ / _   _  _ __  _ __  \ `--.  _   _  _ | |_  ___  
                | ___ \| | | || '__|| '_ \  `--. \| | | || || __|/ _ \ 
                | |_/ /| |_| || |   | |_\) |/\__/ /| |_| || || |_|  __/ 
                \____/  \__,_||_|   | .__/ \____/  \__,_||_| \__|\___| 
                                    | |                                
                                    |_|                                
                       ______                                          
                       | ___ \                                         
                       | |_/ /__ _  _ __  ___   ___  _ __              
                       |  __// _` || '__|/ __| / _ \| '__|             
                       | |  | (_| || |   \__ \|  __/| |                
                       \_|   \__,_||_|   |___/ \___||_|                
                                                           
The Burp Suite XML Scan Parser is a command-line tool designed to parse, analyze, and 
display vulnerability scan results from Burp Suite Professional's XML export format. 
It provides an interactive interface for security professionals to review and analyze 
security findings in a structured and readable format.          

Created by: Aswin Gopalakrishnan
Linkedin: https://www.linkedin.com/in/aswingopalakrishnan/

        """)
        print("\nBURP SUITE SCAN FINDINGS\n")
        burp_parser.display_table()
        
        print("\nOptions:")
        print("1-N: View finding details")
        print("d: Delete a finding")
        print("m: Modify a finding")
        print("s: Save to JSON")
        print("q: Quit")
        
        finding_numbers = []
        choice = input("\nEnter your choice: ").lower()
        
        if choice == 'q':
            break
        elif choice == 'd':
            finding_numbers = str(input("Enter finding number to delete (comma-separated values accepted): "))
            if burp_parser.delete_finding(finding_numbers):
                input("\nPress Enter to continue...")
        elif choice == 'm':
            finding_num = int(input("Enter finding number to modify: "))
            burp_parser.modify_finding(finding_num)
            input("\nPress Enter to continue...")
        elif choice == 's':
            content = burp_parser.save_to_json()
            generator = SecurityReportGenerator()
            generator.make_report(content)
            input("\nPress Enter to continue...")
        else:
            try:
                finding_number = int(choice)
                burp_parser.show_detailed_finding(finding_number)
                input("\nPress Enter to continue...")
            except ValueError:
                print("Please enter a valid option.")
                input("Press Enter to continue...")
        
except ET.ParseError:
    print("Error: Invalid XML file. Please ensure it's a valid Burp Suite export.")
except FileNotFoundError:
    print(f"Error: File not found - {args.xml_file}")

